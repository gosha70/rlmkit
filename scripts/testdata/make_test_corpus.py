"""Generate the deterministic test corpus used by the manual release test plan.

The plan (`doc_internal/<version>/MANUAL_TEST_PLAN.md` §0) asks for a fixed set
of documents so that runs are comparable between the release-candidate pass and
the release pass, and between providers. Hand-assembling those every time is
slow and makes results non-reproducible, so this builds them from a fixed seed.

What it writes (default: ``~/rlm-studio-testdata/``):

===============  ==========================================================
``small.md``     ~5K tokens — Direct-mode path, and Auto's lower branch
``large.md``     ~50K tokens — the core RLM scenario
``huge.md``      ~150K tokens — budget / timeout behaviour, prefix caching
``needle.md``    ~50K tokens with ONE planted fact — the correctness check
``report.pdf``   small text PDF — exercises pypdf extraction on upload
``spec.docx``    small DOCX — exercises python-docx extraction on upload
``ANSWERS.md``   the expected answers, so a run can be graded without guessing
===============  ==========================================================

Generation is seeded, so re-running reproduces the markdown documents and the
PDF byte-for-byte — the script prints a checksum over them. ``spec.docx`` is the
one exception: python-docx stamps creation times into the document and the zip,
so its bytes change every run and it is excluded from the checksum. Sizes are
measured with ``tiktoken`` when available (it ships with litellm) and fall back
to a 4-chars-per-token estimate otherwise.

Usage::

    uv run python scripts/testdata/make_test_corpus.py
    uv run python scripts/testdata/make_test_corpus.py --out /tmp/corpus --force
"""

from __future__ import annotations

import argparse
import random
import sys
import zlib
from pathlib import Path

# The planted fact. Deliberately unguessable — a model that has not read the
# document cannot produce it, which is what makes the needle test meaningful.
NEEDLE_QUESTION = "What is the designated maintenance window for the Orion cluster?"
NEEDLE_ANSWER = "03:15 UTC on the second Tuesday of each month"
NEEDLE_SENTENCE = (
    "The designated maintenance window for the Orion cluster is "
    "03:15 UTC on the second Tuesday of each month."
)

# A second fact, buried in `large.md`, for a non-needle comprehension check.
LARGE_QUESTION = "Which team owns the Tessellate service?"
LARGE_ANSWER = "the Harbour Platform team"
LARGE_SENTENCE = "Ownership of the Tessellate service sits with the Harbour Platform team."

_TOPICS = [
    "ingest pipeline",
    "retention policy",
    "on-call rotation",
    "schema migration",
    "cache invalidation",
    "rate limiting",
    "backfill job",
    "index compaction",
    "quota accounting",
    "shadow traffic",
]
_VERBS = [
    "records",
    "reconciles",
    "drains",
    "replays",
    "throttles",
    "partitions",
    "checkpoints",
    "rebalances",
]
_NOUNS = [
    "the write-ahead log",
    "the staging bucket",
    "the consumer group",
    "the replica set",
    "the audit trail",
    "the dead-letter queue",
    "the shard map",
    "the retry budget",
]


def _sentence(rng: random.Random) -> str:
    return (
        f"The {rng.choice(_TOPICS)} {rng.choice(_VERBS)} {rng.choice(_NOUNS)} "
        f"every {rng.randint(2, 90)} minutes."
    )


def _document(rng: random.Random, sections: int, *, title: str) -> list[str]:
    """Build a plausible technical document as a list of markdown blocks."""
    blocks = [f"# {title}\n"]
    for i in range(sections):
        blocks.append(f"## Section {i}: {_TOPICS[i % len(_TOPICS)].title()}\n")
        for _ in range(rng.randint(3, 6)):
            para = " ".join(_sentence(rng) for _ in range(rng.randint(4, 9)))
            blocks.append(para + "\n")
    return blocks


def _count_tokens(text: str) -> tuple[int, str]:
    """Return (tokens, how). Prefers tiktoken; falls back to a chars/4 estimate."""
    try:
        import tiktoken

        enc = tiktoken.get_encoding("cl100k_base")
        return len(enc.encode(text)), "tiktoken"
    except Exception:
        return len(text) // 4, "estimate (chars/4)"


def _grow_to(rng: random.Random, blocks: list[str], target_tokens: int, *, title: str) -> str:
    """Append sections until the document reaches roughly ``target_tokens``."""
    text = "\n".join(blocks)
    i = len(blocks)
    while _count_tokens(text)[0] < target_tokens:
        extra = _document(rng, 6, title=title)[1:]  # skip the duplicated H1
        blocks.extend(f"{b}" for b in extra)
        i += 1
        text = "\n".join(blocks)
    return text


def _insert_midway(text: str, sentence: str) -> str:
    """Plant a sentence near the middle of the document, inside a paragraph.

    Middle placement matters: at the top a Direct-mode truncation would still
    see it, and at the very end a naive tail-read would. The needle should only
    be findable by something that actually navigates the document.
    """
    paragraphs = text.split("\n\n")
    mid = len(paragraphs) // 2
    paragraphs[mid] = paragraphs[mid].rstrip() + " " + sentence
    return "\n\n".join(paragraphs)


def _minimal_pdf(lines: list[str]) -> bytes:
    """Build a small, valid, text-extractable PDF without a PDF dependency.

    The app extracts PDF text with pypdf, so the file has to be genuinely
    parseable — not a stub. This writes an uncompressed content stream with
    Tj-drawn lines, plus a correct xref table.
    """
    esc = [line.replace("\\", r"\\").replace("(", r"\(").replace(")", r"\)") for line in lines]
    body = "BT\n/F1 12 Tf\n72 720 Td\n14 TL\n" + "".join(f"({t}) Tj\nT*\n" for t in esc) + "ET"
    stream = body.encode("latin-1")

    objects: list[bytes] = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
        b"/Resources << /Font << /F1 5 0 R >> >> /Contents 4 0 R >>",
        b"<< /Length " + str(len(stream)).encode() + b" >>\nstream\n" + stream + b"\nendstream",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    ]

    out = bytearray(b"%PDF-1.4\n")
    offsets: list[int] = []
    for num, obj in enumerate(objects, start=1):
        offsets.append(len(out))
        out += f"{num} 0 obj\n".encode() + obj + b"\nendobj\n"

    xref_at = len(out)
    out += f"xref\n0 {len(objects) + 1}\n".encode()
    out += b"0000000000 65535 f \n"
    for off in offsets:
        out += f"{off:010d} 00000 n \n".encode()
    out += (
        f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\nstartxref\n{xref_at}\n%%EOF\n"
    ).encode()
    return bytes(out)


def _write_docx(path: Path, title: str, paragraphs: list[str]) -> bool:
    try:
        from docx import Document  # type: ignore[import-not-found]
    except ImportError:
        return False
    doc = Document()
    doc.add_heading(title, level=1)
    for para in paragraphs:
        doc.add_paragraph(para)
    doc.save(str(path))
    return True


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(
        description="Generate the deterministic manual-test corpus for RLM Studio."
    )
    parser.add_argument(
        "--out",
        type=Path,
        default=Path.home() / "rlm-studio-testdata",
        help="Output directory (default: ~/rlm-studio-testdata)",
    )
    parser.add_argument(
        "--force", action="store_true", help="Overwrite existing files instead of skipping."
    )
    parser.add_argument("--seed", type=int, default=20260824, help="RNG seed (default: 20260824)")
    args = parser.parse_args(argv)

    out: Path = args.out.expanduser()
    out.mkdir(parents=True, exist_ok=True)

    written: list[tuple[str, str, str]] = []

    def emit(name: str, data: str | bytes) -> None:
        path = out / name
        if path.exists() and not args.force:
            print(f"  skip   {name} (exists — pass --force to overwrite)")
            return
        if isinstance(data, bytes):
            path.write_bytes(data)
            size = f"{len(data) / 1024:.1f} KiB"
            written.append((name, size, "binary"))
        else:
            path.write_text(data)
            tokens, how = _count_tokens(data)
            written.append((name, f"{tokens:,} tokens", how))
        print(f"  write  {name}")

    # --- markdown documents -------------------------------------------------
    rng = random.Random(args.seed)
    small = _grow_to(
        rng, _document(rng, 4, title="Platform Handbook"), 5_000, title="Platform Handbook"
    )
    emit("small.md", small)

    rng = random.Random(args.seed + 1)
    large_blocks = _document(rng, 20, title="Service Catalogue")
    large = _grow_to(rng, large_blocks, 50_000, title="Service Catalogue")
    large = _insert_midway(large, LARGE_SENTENCE)
    emit("large.md", large)

    rng = random.Random(args.seed + 2)
    huge = _grow_to(
        rng,
        _document(rng, 40, title="Operations Compendium"),
        150_000,
        title="Operations Compendium",
    )
    emit("huge.md", huge)

    rng = random.Random(args.seed + 3)
    needle = _grow_to(
        rng, _document(rng, 20, title="Runbook Archive"), 50_000, title="Runbook Archive"
    )
    needle = _insert_midway(needle, NEEDLE_SENTENCE)
    emit("needle.md", needle)

    # --- binary formats -----------------------------------------------------
    emit(
        "report.pdf",
        _minimal_pdf(
            [
                "Quarterly Reliability Report",
                "",
                "The Orion cluster completed 41 scheduled failovers this quarter.",
                "Mean time to recovery was 6.4 minutes across all regions.",
                "The Tessellate service reported no Sev-1 incidents.",
            ]
        ),
    )

    docx_ok = _write_docx(
        out / "spec.docx",
        "Interface Specification",
        [
            "This document specifies the ingest contract for the Tessellate service.",
            "Requests are accepted as newline-delimited JSON, at most 8 MiB per batch.",
            "The service acknowledges a batch only after the write-ahead log is durable.",
            "Retries use the shared retry budget and are capped at four attempts.",
        ],
    )
    if docx_ok:
        written.append(
            ("spec.docx", f"{(out / 'spec.docx').stat().st_size / 1024:.1f} KiB", "binary")
        )
        print("  write  spec.docx")
    else:
        print("  SKIP   spec.docx — python-docx not installed (uv sync --extra studio)")

    # --- answer key ---------------------------------------------------------
    emit(
        "ANSWERS.md",
        (
            "# Expected answers\n\n"
            "Generated by `scripts/testdata/make_test_corpus.py` "
            f"(seed {args.seed}). Regenerating with the same seed reproduces these exactly.\n\n"
            f"## needle.md\n\n**Q:** {NEEDLE_QUESTION}\n\n**A:** {NEEDLE_ANSWER}\n\n"
            "The sentence is planted mid-document on purpose: a model that only reads the\n"
            "head or the tail cannot answer it, so this distinguishes real navigation from\n"
            "truncation luck. Grade RLM and Direct separately and record both.\n\n"
            f"## large.md\n\n**Q:** {LARGE_QUESTION}\n\n**A:** {LARGE_ANSWER}\n\n"
            "## report.pdf\n\n**Q:** How many scheduled failovers did the Orion cluster "
            "complete?\n\n**A:** 41\n\n"
            "## spec.docx\n\n**Q:** What is the maximum batch size for the ingest "
            "contract?\n\n**A:** 8 MiB\n"
        ),
    )

    print(f"\nCorpus at: {out}")
    if written:
        width = max(len(n) for n, _, _ in written)
        for name, size, how in written:
            print(f"  {name:<{width}}  {size:>16}   {how}")
    # Checksum only the generated text + PDF. `spec.docx` is deliberately
    # excluded: python-docx stamps creation/modification times into the
    # document and into the zip entries, so its bytes change on every run even
    # with a fixed seed. The documents that decide whether two test runs are
    # comparable are byte-identical, and that is what this covers.
    deterministic = sorted(p for p in out.iterdir() if p.is_file() and p.suffix in {".md", ".pdf"})
    checksum = zlib.crc32(b"".join(p.read_bytes() for p in deterministic))
    print(
        f"\ncorpus checksum: {checksum:08x}  "
        f"({len(deterministic)} deterministic files; same seed → same checksum)"
    )
    print("note: spec.docx carries generation timestamps and is not part of the checksum.")
    return 0


if __name__ == "__main__":
    sys.exit(main())
